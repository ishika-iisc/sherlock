from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
LATEX_DIR = ROOT
WORD_TEMPLATE = Path("/Users/abhi/Downloads/templates/word/project-report-template.docx")
SEAL_IMAGE = Path("/Users/abhi/Downloads/templates/latex/IISc_Master_Seal_Black.jpg")
OUTPUT = ROOT / "mtech_thesis_document_intelligence.docx"
ASSETS = ROOT / "word_assets"
BULLET_NUM_ID: int | None = None
DECIMAL_NUM_ID: int | None = None
BULLET_ABS_ID: int | None = None
DECIMAL_ABS_ID: int | None = None

TITLE = "AI-Powered Document Intelligence System for Automated Extraction, Validation, Search, and Question Answering"
AUTHOR = "Student Name"
SRNO = "SR Number"
SUBMIT_DATE = "Month, Year"
STREAM = "Artificial Intelligence"
ORGANISATION = "Organisation Name"
GUIDE = "Project Guide Name"
MENTOR = "IISc Faculty Mentor Name"
MENTOR_DEPT = "Faculty Mentor Department Name"
DEGREE_START = "Join Month, Year"
DEGREE_END = "End Month, Year"

CITATION_ORDER = [
    "smith2007tesseract",
    "easyocr2026",
    "devlin2019bert",
    "reimers2019sentencebert",
    "huang2022layoutlmv3",
    "lewis2020rag",
    "react2026",
    "vite2026",
    "fastapi2026",
    "sqlalchemy2026",
]
CITATION_NUM = {key: idx + 1 for idx, key in enumerate(CITATION_ORDER)}

CREFS = {
    "chap:literature_review": "Chapter 2",
    "chap:methodology": "Chapter 3",
    "chap:experiments_and_results": "Chapter 4",
    "chap:conclusion": "Chapter 5",
    "fig:architecture": "Figure 1",
    "tab:technology_stack": "Table 1",
    "alg:document_pipeline": "Algorithm 1",
    "alg:agentic_rag": "Algorithm 2",
    "tab:evaluation_results": "Table 2",
    "tab:operational_routes": "Table 3",
}


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color: str = "BFBFBF", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def _next_numbering_ids(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = []
    num_ids = []
    for abstract in numbering.findall(qn("w:abstractNum")):
        value = abstract.get(qn("w:abstractNumId"))
        if value and value.isdigit():
            abstract_ids.append(int(value))
    for num in numbering.findall(qn("w:num")):
        value = num.get(qn("w:numId"))
        if value and value.isdigit():
            num_ids.append(int(value))
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def _make_abstract_num(abstract_id: int, kind: str) -> OxmlElement:
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)

    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Symbol" if kind == "bullet" else "Calibri")
    rfonts.set(qn("w:hAnsi"), "Symbol" if kind == "bullet" else "Calibri")
    rpr.append(rfonts)
    lvl.append(rpr)

    abstract.append(lvl)
    return abstract


def _make_num(num_id: int, abstract_id: int) -> OxmlElement:
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    return num


def setup_numbering(doc: Document) -> None:
    global BULLET_NUM_ID, DECIMAL_NUM_ID, BULLET_ABS_ID, DECIMAL_ABS_ID
    numbering = doc.part.numbering_part.element
    bullet_abs, bullet_num = _next_numbering_ids(doc)
    decimal_abs = bullet_abs + 1
    decimal_num = bullet_num + 1
    numbering.append(_make_abstract_num(bullet_abs, "bullet"))
    numbering.append(_make_num(bullet_num, bullet_abs))
    numbering.append(_make_abstract_num(decimal_abs, "decimal"))
    numbering.append(_make_num(decimal_num, decimal_abs))
    BULLET_ABS_ID = bullet_abs
    DECIMAL_ABS_ID = decimal_abs
    BULLET_NUM_ID = bullet_num
    DECIMAL_NUM_ID = decimal_num


def create_numbering_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    _, next_num = _next_numbering_ids(doc)
    numbering.append(_make_num(next_num, abstract_id))
    return next_num


def new_decimal_list(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = _next_numbering_ids(doc)
    numbering.append(_make_abstract_num(abstract_id, "decimal"))
    numbering.append(_make_num(num_id, abstract_id))
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = numpr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        numpr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = numpr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        numpr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field in Microsoft Word to generate the table of contents."
    fld_sep.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    body = doc.styles["Body Text"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AI-Powered Document Intelligence Thesis Draft")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, color=None, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_body(doc: Document, text: str, style: str = "Body Text"):
    if not text:
        return
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    if BULLET_NUM_ID is not None:
        apply_numbering(p, BULLET_NUM_ID)
    p.add_run(text)


def add_numbered(doc: Document, text: str, num_id: int | None = None):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    active_num_id = num_id if num_id is not None else DECIMAL_NUM_ID
    if active_num_id is not None:
        apply_numbering(p, active_num_id)
    p.add_run(text)


def add_cover(doc: Document) -> None:
    add_centered(doc, TITLE, size=18, bold=True, color=(0x0B, 0x25, 0x45))
    doc.add_paragraph()
    add_centered(doc, "A PROJECT REPORT", size=12, bold=True)
    add_centered(doc, "SUBMITTED IN PARTIAL FULFILMENT OF THE", size=12, bold=True)
    add_centered(doc, "REQUIREMENTS FOR THE DEGREE OF", size=12, bold=True)
    add_centered(doc, "Master of Technology (Online)", size=14, bold=True, color=(0x2E, 0x74, 0xB5))
    add_centered(doc, "IN", size=12, bold=True)
    add_centered(doc, STREAM, size=14, bold=True, color=(0x2E, 0x74, 0xB5))
    doc.add_paragraph()
    add_centered(doc, "BY", size=12, bold=True)
    add_centered(doc, AUTHOR, size=14, bold=True)
    add_centered(doc, f"SR No. {SRNO}", size=11)
    doc.add_paragraph()
    if SEAL_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(SEAL_IMAGE), width=Inches(1.25))
    doc.add_paragraph()
    add_centered(doc, "Faculty of Engineering", size=12, bold=True)
    add_centered(doc, "Indian Institute of Science", size=12, bold=True)
    add_centered(doc, "Bangalore - 560012 (INDIA)", size=12, bold=True)
    doc.add_paragraph()
    add_centered(doc, SUBMIT_DATE, size=12, bold=True)
    doc.add_page_break()


def add_declaration(doc: Document) -> None:
    p = doc.add_heading("Declaration of Originality", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(
        doc,
        f"I, {AUTHOR}, with SR No. {SRNO} hereby declare that the material presented in the project report titled",
    )
    add_centered(doc, TITLE, size=12, bold=True)
    add_body(
        doc,
        f"represents original work carried out by me at {ORGANISATION} as part of the project credit requirements of the Master of Technology (Online) degree in {STREAM} at the Indian Institute of Science, between {DEGREE_START} to {DEGREE_END}.",
    )
    add_body(doc, "With my signature, I certify that:")
    for item in [
        "I have not manipulated any of the data or results.",
        "I have not committed any plagiarism of intellectual property. I have clearly indicated and referenced the contributions of others.",
        "I have explicitly acknowledged all collaborative research and discussions.",
        "I have understood that any false claim will result in severe disciplinary action.",
        "I have understood that the work may be screened for any form of academic misconduct.",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()
    add_body(doc, "Date: ____________________________                                      Student Signature: ____________________________")
    doc.add_paragraph()
    add_body(
        doc,
        "In our capacities as internal project guide and faculty mentor of the above-mentioned work, we certify that the above statements are true to the best of our knowledge, and we have carried out due diligence to ensure the originality of the report.",
    )
    doc.add_paragraph()
    add_body(doc, f"Internal Guide Name: {GUIDE}                                      Internal Guide Signature: ____________________________")
    add_body(doc, f"Organisation: {ORGANISATION}")
    add_body(doc, f"Faculty Mentor Name: {MENTOR}                                      Faculty Mentor Signature: ____________________________")
    add_body(doc, f"Department: {MENTOR_DEPT}")
    doc.add_page_break()


def add_copyright(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    add_centered(doc, f"© {AUTHOR}", size=13, bold=True)
    add_centered(doc, SUBMIT_DATE, size=13, bold=True)
    add_centered(doc, "All rights reserved", size=13, bold=True)
    doc.add_page_break()


def clean_latex(text: str) -> str:
    text = text.strip()
    text = text.replace("``", '"').replace("''", '"')
    text = text.replace("\\&", "&").replace("\\_", "_").replace("---", "-")
    text = re.sub(r"\\cref\{([^}]+)\}", lambda m: CREFS.get(m.group(1), m.group(1)), text)
    text = re.sub(r"~?\\cite\{([^}]+)\}", lambda m: "[" + ", ".join(str(CITATION_NUM.get(k.strip(), "?")) for k in m.group(1).split(",")) + "]", text)
    text = re.sub(r"\\texttt\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\paragraph\{\}\s*", "", text)
    text = text.replace("\\(", "").replace("\\)", "")
    text = text.replace("$", "")
    text = text.replace("\\theta", "theta").replace("\\mid", "|")
    text = text.replace("\\textbackslash", "\\")
    text = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?(\{[^}]*\})?", "", text)
    text = text.replace("{", "").replace("}", "")
    return " ".join(text.split())


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def parse_table(block: str) -> tuple[str, list[list[str]]]:
    cap = re.search(r"\\caption\{([^}]*)\}", block, re.S)
    caption = clean_latex(cap.group(1)) if cap else "Table"
    rows = []
    for raw in block.splitlines():
        if "&" not in raw or "\\\\" not in raw:
            continue
        raw = raw.split("\\\\")[0]
        raw = raw.replace("\\bf", "")
        cells = [clean_latex(cell) for cell in raw.split("&")]
        if cells:
            rows.append(cells)
    return caption, rows


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    if "Normal Table" in [style.name for style in doc.styles]:
        table.style = "Normal Table"
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(table.cell(r_idx, c_idx), "F2F4F7")
    if cols == 2:
        widths = [3000, 6360]
    elif cols == 3:
        widths = [3600, 2600, 3160]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    add_caption(doc, caption)


def make_architecture_image(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34)
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 24)
    except Exception:
        font_title = ImageFont.load_default()
        font = ImageFont.load_default()

    def box(x, y, w, h, text, fill="#F2F4F7"):
        draw.rounded_rectangle([x, y, x + w, y + h], radius=18, fill=fill, outline="#2E74B5", width=3)
        lines = text.split("\n")
        total_h = len(lines) * 30
        for i, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text((x + (w - (bbox[2] - bbox[0])) / 2, y + (h - total_h) / 2 + i * 32), line, fill="#0B2545", font=font)

    def arrow(x1, y1, x2, y2):
        draw.line([x1, y1, x2, y2], fill="#0B2545", width=4)
        draw.polygon([(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)], fill="#0B2545")

    draw.text((350, 40), "Document Intelligence System Architecture", fill="#0B2545", font=font_title)
    box(520, 120, 560, 95, "React Frontend\nDashboard, Upload, Search, RAG")
    box(520, 260, 560, 95, "FastAPI Backend\nREST API")
    box(520, 400, 560, 95, "Document Processing Pipeline")
    box(130, 590, 330, 95, "OCR\nTesseract, EasyOCR", "#FFFFFF")
    box(635, 590, 330, 95, "Layout-Aware\nVLM", "#FFFFFF")
    box(1120, 590, 330, 95, "Entity Extraction\nRules and LLM", "#FFFFFF")
    box(520, 760, 560, 95, "SQLite Storage\nDocuments, Extractions, Logs")
    box(1120, 760, 330, 95, "Search and RAG\nEvidence, Answer", "#FFFFFF")
    arrow(800, 215, 800, 260)
    arrow(800, 355, 800, 400)
    arrow(650, 495, 300, 590)
    arrow(800, 495, 800, 590)
    arrow(950, 495, 1280, 590)
    arrow(300, 685, 650, 760)
    arrow(800, 685, 800, 760)
    arrow(1280, 685, 1000, 760)
    arrow(1080, 805, 1120, 805)
    img.save(path)


def parse_algorithm(block: str) -> tuple[str, list[str]]:
    cap = re.search(r"\\caption\{([^}]*)\}", block, re.S)
    caption = clean_latex(cap.group(1)) if cap else "Algorithm"
    steps = []
    for line in block.splitlines():
        line = line.strip()
        if line.startswith("\\State"):
            steps.append(clean_latex(line.replace("\\State", "", 1)))
        elif line.startswith("\\If"):
            steps.append("If " + clean_latex(line.replace("\\If", "", 1)))
        elif line.startswith("\\Else"):
            steps.append("Else")
        elif line.startswith("\\EndIf"):
            steps.append("End if")
    return caption, steps


def process_block(doc: Document, block: str) -> None:
    block = block.strip()
    if not block:
        return
    if block.startswith("\\begin{table}"):
        caption, rows = parse_table(block)
        add_table(doc, caption, rows)
        return
    if block.startswith("\\begin{algorithm}"):
        caption, steps = parse_algorithm(block)
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
        num_id = new_decimal_list(doc)
        for step in steps:
            add_numbered(doc, step, num_id=num_id)
        return
    if block.startswith("\\begin{figure}"):
        if "architecture" in block.lower():
            img_path = ASSETS / "architecture.png"
            make_architecture_image(img_path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img_path), width=Inches(6.2))
            add_caption(doc, "Figure 1: High-level architecture of the document intelligence system")
        return
    if block.startswith("\\begin{quote}"):
        text = block.replace("\\begin{quote}", "").replace("\\end{quote}", "")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.right_indent = Inches(0.25)
        run = p.add_run(clean_latex(text))
        run.italic = True
        run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
        return
    if block.startswith("\\begin{enumerate}"):
        num_id = new_decimal_list(doc)
        for item in re.findall(r"\\item\s+(.+?)(?=(?:\n\s*\\item)|(?:\n\s*\\end\{enumerate\})|$)", block, re.S):
            add_numbered(doc, clean_latex(item), num_id=num_id)
        return

    for raw_line in block.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        chap = re.match(r"\\chapter\{([^}]*)\}", line)
        sec = re.match(r"\\section\{([^}]*)\}", line)
        sub = re.match(r"\\subsection\{([^}]*)\}", line)
        if chap:
            if len(doc.paragraphs) > 0:
                doc.add_page_break()
            doc.add_heading(clean_latex(chap.group(1)), level=1)
        elif sec:
            doc.add_heading(clean_latex(sec.group(1)), level=2)
        elif sub:
            doc.add_heading(clean_latex(sub.group(1)), level=3)
        elif line.startswith("\\label"):
            continue
        else:
            add_body(doc, clean_latex(line))


def process_tex_file(doc: Document, path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    text = re.sub(r"(?m)^%.*\n?", "", text)
    tokens = re.split(
        r"(\\begin\{table\}.*?\\end\{table\}|\\begin\{algorithm\}.*?\\end\{algorithm\}|\\begin\{figure\}.*?\\end\{figure\}|\\begin\{enumerate\}.*?\\end\{enumerate\}|\\begin\{quote\}.*?\\end\{quote\})",
        text,
        flags=re.S,
    )
    for token in tokens:
        process_block(doc, token)


def add_front_matter_file(doc: Document, title: str, path: Path) -> None:
    doc.add_heading(title, level=1)
    text = path.read_text(encoding="utf-8")
    for para in re.findall(r"\\paragraph\{\}\s*(.*?)(?=\n\\paragraph\{\}|$)", text, flags=re.S):
        add_body(doc, clean_latex(para))
    doc.add_page_break()


def add_toc_page(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    entries = [
        ("Acknowledgements", "4"),
        ("Abstract", "5"),
        ("1. Introduction", "7"),
        ("2. Literature Review", "10"),
        ("3. Methodology and Implementation", "13"),
        ("4. Experiments and Results", "18"),
        ("5. Conclusions and Future Work", "21"),
        ("Appendix A. API Endpoints", "23"),
        ("Appendix B. Prototype Configuration", "24"),
        ("Appendix C. Recommended Screenshots", "25"),
        ("References", "26"),
    ]
    table = doc.add_table(rows=len(entries) + 1, cols=2)
    set_cell_text(table.cell(0, 0), "Section", bold=True)
    set_cell_text(table.cell(0, 1), "Draft page", bold=True)
    set_cell_shading(table.cell(0, 0), "F2F4F7")
    set_cell_shading(table.cell(0, 1), "F2F4F7")
    for idx, (label, page) in enumerate(entries, 1):
        set_cell_text(table.cell(idx, 0), label)
        set_cell_text(table.cell(idx, 1), page)
    set_table_borders(table)
    set_table_geometry(table, [7800, 1560])


def add_references(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("References", level=1)
    refs = [
        "Smith, R. (2007). An Overview of the Tesseract OCR Engine. Proceedings of the Ninth International Conference on Document Analysis and Recognition, 629-633. doi:10.1109/ICDAR.2007.4376991.",
        "JaidedAI. (2026). EasyOCR: Ready-to-use OCR with 80+ supported languages. https://github.com/JaidedAI/EasyOCR. Accessed: 2026-06-08.",
        "Devlin, J., Chang, M.-W., Lee, K., and Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. arXiv:1810.04805.",
        "Reimers, N., and Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. Proceedings of EMNLP 2019, 3982-3992. doi:10.18653/v1/D19-1410.",
        "Huang, Y., Lv, T., Cui, L., Lu, Y., and Wei, F. (2022). LayoutLMv3: Pre-training for Document AI with Unified Text and Image Masking. arXiv:2204.08387.",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Kuttler, H., Lewis, M., Yih, W., Rocktaschel, T., Riedel, S., and Kiela, D. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
        "React Team. (2026). React Documentation. https://react.dev/. Accessed: 2026-06-08.",
        "Vite. (2026). Vite Documentation. https://vite.dev/guide/. Accessed: 2026-06-08.",
        "FastAPI. (2026). FastAPI Documentation. https://fastapi.tiangolo.com/. Accessed: 2026-06-08.",
        "SQLAlchemy. (2026). SQLAlchemy Documentation. https://docs.sqlalchemy.org/. Accessed: 2026-06-08.",
    ]
    num_id = new_decimal_list(doc)
    for ref in refs:
        add_numbered(doc, ref, num_id=num_id)


def ensure_valid_docx(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as zf:
        needed = {"[Content_Types].xml", "word/document.xml"}
        missing = needed - set(zf.namelist())
        if missing:
            raise RuntimeError(f"DOCX missing required parts: {missing}")


def build() -> None:
    ASSETS.mkdir(exist_ok=True)
    doc = Document(str(WORD_TEMPLATE))
    clear_body(doc)
    setup_document(doc)
    setup_numbering(doc)

    add_cover(doc)
    add_declaration(doc)
    add_copyright(doc)
    add_front_matter_file(doc, "Acknowledgements", LATEX_DIR / "acknowledgements.tex")
    add_front_matter_file(doc, "Abstract", LATEX_DIR / "abstract.tex")
    add_toc_page(doc)

    for filename in [
        "chap_introduction.tex",
        "chap_literature_review.tex",
        "chap_methodology.tex",
        "chap_experiments_and_results.tex",
        "chap_conclusion.tex",
        "appendix.tex",
    ]:
        process_tex_file(doc, LATEX_DIR / filename)

    add_references(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.author = AUTHOR
    doc.core_properties.subject = "M.Tech Project Report"
    doc.core_properties.comments = "Generated from the Document Intelligence thesis draft."
    doc.save(str(OUTPUT))
    ensure_valid_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
